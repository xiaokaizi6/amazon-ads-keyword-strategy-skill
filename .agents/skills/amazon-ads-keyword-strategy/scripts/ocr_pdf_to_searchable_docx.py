#!/usr/bin/env python3
"""Create a page-traceable OCR Word derivative from rendered PDF page images.

The original PDF remains authoritative. This script deliberately preserves all
OCR lines (including low-confidence text) in JSONL and emits page headings in
the DOCX, so readers can return to the original page for tables, screenshots,
and uncertain recognition.
"""

from __future__ import annotations

import argparse
import hashlib
import json
import re
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from rapidocr_onnxruntime import RapidOCR


BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
MUTED = "666666"


def sha256(path: Path) -> str:
    digest = hashlib.sha256()
    with path.open("rb") as handle:
        for chunk in iter(lambda: handle.read(1024 * 1024), b""):
            digest.update(chunk)
    return digest.hexdigest()


def set_font(run, *, size: float, bold: bool = False, color: str = "000000") -> None:
    run.font.name = "Microsoft YaHei"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    run._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    run.font.size = Pt(size)
    run.bold = bold
    run.font.color.rgb = RGBColor.from_string(color)


def configure_document(document: Document) -> None:
    section = document.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    normal = document.styles["Normal"]
    normal.font.name = "Microsoft YaHei"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25
    for name, size, color, before, after in [
        ("Heading 1", 16, BLUE, 18, 10),
        ("Heading 2", 13, BLUE, 14, 7),
    ]:
        style = document.styles[name]
        style.font.name = "Microsoft YaHei"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.font.bold = True
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    header_run = header.add_run("PDF OCR 文字检索版（派生物，原 PDF 为准）")
    set_font(header_run, size=8.5, color=MUTED)


def natural_page_key(path: Path) -> int:
    match = re.search(r"(\d+)$", path.stem)
    if not match:
        raise ValueError(f"page number not found in {path.name}")
    return int(match.group(1))


def line_key(item: list[object]) -> tuple[float, float]:
    box = item[0]
    if not isinstance(box, list) or not box:
        return (0.0, 0.0)
    return (float(box[0][1]), float(box[0][0]))


def main() -> None:
    parser = argparse.ArgumentParser()
    parser.add_argument("--source-pdf", type=Path, required=True)
    parser.add_argument("--page-images-dir", type=Path, required=True)
    parser.add_argument("--output-docx", type=Path, required=True)
    parser.add_argument("--output-jsonl", type=Path, required=True)
    parser.add_argument("--source-id", required=True)
    args = parser.parse_args()

    pages = sorted(args.page_images_dir.glob("page-*.png"), key=natural_page_key)
    if not pages:
        raise SystemExit("no page PNGs found")
    expected = list(range(1, len(pages) + 1))
    actual = [natural_page_key(path) for path in pages]
    if actual != expected:
        raise SystemExit(f"non-contiguous page images: {actual[:3]}...{actual[-3:]}")

    args.output_docx.parent.mkdir(parents=True, exist_ok=True)
    args.output_jsonl.parent.mkdir(parents=True, exist_ok=True)
    document = Document()
    configure_document(document)

    title = document.add_paragraph()
    title.paragraph_format.space_before = Pt(12)
    title.paragraph_format.space_after = Pt(5)
    set_font(title.add_run("亚马逊专题课：进阶广告诊断与优化全指导"), size=22, bold=True, color=DARK_BLUE)
    subtitle = document.add_paragraph()
    subtitle.paragraph_format.space_after = Pt(12)
    set_font(subtitle.add_run("PDF OCR 文字检索版"), size=16, bold=True, color=BLUE)
    notice = document.add_paragraph()
    notice.paragraph_format.space_after = Pt(12)
    set_font(
        notice.add_run(
            f"来源：{args.source_id}；原始 PDF 共 {len(pages)} 页，SHA-256 {sha256(args.source_pdf)}。"
            "本文件由逐页机器 OCR 生成，保留原 PDF 页码以便回查。表格、截图、公式、数字、专有名词和低置信度文字必须以原 PDF 页面为准。"
        ),
        size=10,
        color=MUTED,
    )

    engine = RapidOCR()
    total_lines = 0
    with args.output_jsonl.open("w", encoding="utf-8", newline="\n") as jsonl:
        for page_path in pages:
            page_number = natural_page_key(page_path)
            result, elapsed = engine(str(page_path))
            entries = sorted(result or [], key=line_key)
            lines = [
                {
                    "text": str(entry[1]),
                    "confidence": round(float(entry[2]), 6),
                    "bbox": entry[0],
                }
                for entry in entries
                if len(entry) >= 3 and str(entry[1]).strip()
            ]
            row = {
                "source_id": args.source_id,
                "source_location": f"PDF p.{page_number}",
                "page_number": page_number,
                "page_image": page_path.name,
                "ocr_engine": "rapidocr-onnxruntime",
                "ocr_lines": lines,
                "ocr_text": "\n".join(line["text"] for line in lines),
                "mean_confidence": round(sum(line["confidence"] for line in lines) / len(lines), 6) if lines else None,
                "ocr_elapsed_seconds": [round(float(value), 6) for value in elapsed],
                "review_status": "machine_ocr_derived_not_source_authority",
            }
            jsonl.write(json.dumps(row, ensure_ascii=False, separators=(",", ":")) + "\n")
            total_lines += len(lines)

            heading = document.add_heading(f"原 PDF 第 {page_number} 页", level=2)
            heading.paragraph_format.keep_with_next = True
            if not lines:
                paragraph = document.add_paragraph()
                set_font(paragraph.add_run("[OCR 未识别出文字；请直接查看原 PDF 页面。]"), size=10, color=MUTED)
                continue
            for line in lines:
                paragraph = document.add_paragraph()
                paragraph.paragraph_format.space_after = Pt(3)
                paragraph.paragraph_format.line_spacing = 1.15
                run = paragraph.add_run(line["text"])
                set_font(run, size=10.5, color="000000" if line["confidence"] >= 0.8 else "7A5A00")

    document.save(args.output_docx)
    print(f"OCR pages: {len(pages)}")
    print(f"OCR lines: {total_lines}")
    print(f"DOCX: {args.output_docx}")
    print(f"JSONL: {args.output_jsonl}")


if __name__ == "__main__":
    main()
